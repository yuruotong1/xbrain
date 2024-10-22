import unittest
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from xbrain.utils.upload_util import process_file

class TestHandleDocx(unittest.TestCase):
    def setUp(self):
        # Create a persistent DOCX file in the test folder
        os.mkdir(os.path.join(os.path.dirname(__file__), "temp"))
        self.test_file_path = os.path.join(os.path.dirname(__file__), "temp/test_docx_persistent.docx")
        doc = Document()
        # Set up page properties
        section = doc.sections[0]
        section.page_width = Pt(595.3)  # A4 width in points
        section.page_height = Pt(841.9)  # A4 height in points
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)
        # Add paragraph with different styles
        p = doc.add_paragraph("This is a test paragraph.")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run = p.add_run(" This is bold text.")
        run.bold = True
        run = p.add_run(" This is italic text.")
        run.italic = True
        run = p.add_run(" This is underlined text.")
        run.underline = True
        run = p.add_run(" This is text with a specific font size.")
        run.font.size = Pt(24)
        doc.save(self.test_file_path)

    def test_handle_docx(self):
        # Call the function and check if it returns the expected result
        result = process_file(self.test_file_path)
        print(result)
        self.assertIn("This is a test paragraph.", result)
        self.assertIn("paragraphs", result)
        self.assertIn("This is bold text.", result)
        self.assertIn("bold", result)
        self.assertIn("This is italic text.", result)
        self.assertIn("italic", result)
        self.assertIn("This is underlined text.", result)
        self.assertIn("underline", result)
        self.assertIn("This is text with a specific font size.", result)
        self.assertIn("font_size", result)
        # Check for page size, margins, and line spacing
        self.assertIn("page_size", result)
        self.assertIn("line_spacing", result)
        self.assertIn("top_margin", result)
        self.assertIn("bottom_margin", result)
        self.assertIn("left_margin", result)
        self.assertIn("right_margin", result)

if __name__ == "__main__":
    unittest.main()
